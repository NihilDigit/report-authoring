#!/usr/bin/env python3
"""
inspect_template.py — 阶段 1：摘要 docx 模板的结构与样式。

输出：
  1) 顶层段落 + 表格（递归展开嵌套 cell/table）
  2) sectPr 页面几何
  3) ★ 段落分类：封面 / TOC / 标题 / 正文 / 空段
     —— 封面段标为「保留区」：很多模板用表单字段/文本框/绝对定位魔法，动一动就坏
  4) ★ 样式 diff 表：「正文首段实际样式 vs skill 硬编码默认」
     从分类为 body 的**第一段**取值，不再误取封面右对齐的元信息
     Agent 看到 diff 就知道要不要给 builders.py 传 override

通用性：**不假设模板有表格**（论文模板常是纯段落流）。表格内的段落也会被递归扫到。

用法：
    uvx --with python-docx python inspect_template.py <docx> [--verbose]
"""
from __future__ import annotations

import re
import sys
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET
from typing import Iterator


# -------------- skill 硬编码默认（来自 builders.py，同步更新）--------------
# 注意：这些是「常见约定」而非权威值。见 docs/typography.md 的「诚实」段。
SKILL_DEFAULTS = {
    "body.size":       {"val": 21,  "note": "五号 10.5pt"},
    "body.line":       {"val": 360, "note": "1.5 倍行距"},
    "body.lineRule":   {"val": "auto", "note": ""},
    "body.firstLine":  {"val": 200, "note": "首行缩进 2 字符（firstLineChars）"},
    "body.jc":         {"val": "left", "note": "左对齐（中文正文惯例）"},
    "heading1.size":   {"val": 24,  "note": "小三"},
    "heading2.size":   {"val": 22,  "note": "四号"},
    "heading3.size":   {"val": 21,  "note": "小四"},
    "caption.size":    {"val": 18,  "note": "小五 + 加粗 + 居中"},
    "code.size":       {"val": 20,  "note": "稍小于正文（10pt）"},
    "code.font":       {"val": "Consolas", "note": ""},
}

# 封面 / 元信息关键词：段落含这些字段通常是表单页，不是正文
_COVER_FIELD_KEYWORDS = (
    "课程号", "课程名", "班级", "学号", "姓名", "教师", "指导老师", "指导教师",
    "学院", "专业", "年级", "班号", "组号", "小组",
    "日期", "报告日期", "实验日期", "提交日期", "完成日期",
    "论文题目", "题目", "Title",
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"


# -------------- python-docx 遍历（段落 + 递归表格） --------------

def dump_paragraphs(paragraphs, indent: str = ""):
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        marker = "（空）" if not text else text[:80] + ("…" if len(text) > 80 else "")
        print(f"{indent}[P {i:3d}] {marker}")


def dump_table(tbl, tid: str, indent: str = ""):
    rows, cols = len(tbl.rows), len(tbl.columns)
    print(f"\n{indent}[T {tid}] {rows} rows × {cols} cols")
    for r, row in enumerate(tbl.rows):
        for c, cell in enumerate(row.cells):
            text = cell.text.strip()
            marker = "（空）" if not text else text[:80] + ("…" if len(text) > 80 else "")
            print(f"{indent}  [T {tid} {r},{c}] {marker}")
            if cell.tables:
                for ntid, ntbl in enumerate(cell.tables):
                    dump_table(ntbl, f"{tid}.{r}.{c}.{ntid}", indent + "    ")


# -------------- XML 工具 --------------

def _read_xml(zf: ZipFile, name: str) -> ET.Element | None:
    try:
        with zf.open(name) as f:
            return ET.fromstring(f.read())
    except KeyError:
        return None


def _get(elem, tag: str):
    return elem.find(f"{W}{tag}") if elem is not None else None


def _attr(node, attr: str) -> str | None:
    if node is None:
        return None
    return node.get(f"{W}{attr}")


# -------------- styles.xml 提取默认 --------------

def extract_default_styles(styles_root: ET.Element | None) -> dict:
    out = {}
    if styles_root is None:
        return out

    rpr_default = styles_root.find(f"{W}docDefaults/{W}rPrDefault/{W}rPr")
    if rpr_default is not None:
        sz = _get(rpr_default, "sz")
        out["docDefault.size"] = _attr(sz, "val")
        fonts = _get(rpr_default, "rFonts")
        if fonts is not None:
            out["docDefault.ascii"] = fonts.get(f"{W}ascii")
            out["docDefault.eastAsia"] = fonts.get(f"{W}eastAsia")

    ppr_default = styles_root.find(f"{W}docDefaults/{W}pPrDefault/{W}pPr")
    if ppr_default is not None:
        spacing = _get(ppr_default, "spacing")
        if spacing is not None:
            out["docDefault.line"] = spacing.get(f"{W}line")
            out["docDefault.lineRule"] = spacing.get(f"{W}lineRule")
        ind = _get(ppr_default, "ind")
        if ind is not None:
            out["docDefault.firstLineChars"] = ind.get(f"{W}firstLineChars")

    for style in styles_root.findall(f"{W}style"):
        if style.get(f"{W}type") == "paragraph" and style.get(f"{W}default") == "1":
            sid = style.get(f"{W}styleId", "?")
            out["normal.styleId"] = sid
            rpr = _get(style, "rPr")
            if rpr is not None:
                sz = _get(rpr, "sz")
                if sz is not None:
                    out["normal.size"] = sz.get(f"{W}val")
                fonts = _get(rpr, "rFonts")
                if fonts is not None:
                    out["normal.ascii"] = fonts.get(f"{W}ascii")
                    out["normal.eastAsia"] = fonts.get(f"{W}eastAsia")
            ppr = _get(style, "pPr")
            if ppr is not None:
                spacing = _get(ppr, "spacing")
                if spacing is not None:
                    out["normal.line"] = spacing.get(f"{W}line")
                    out["normal.lineRule"] = spacing.get(f"{W}lineRule")
                ind = _get(ppr, "ind")
                if ind is not None:
                    out["normal.firstLineChars"] = ind.get(f"{W}firstLineChars")
                jc = _get(ppr, "jc")
                if jc is not None:
                    out["normal.jc"] = jc.get(f"{W}val")
            break
    return out


# -------------- 通用段落遍历（不假设有表格） --------------

def iter_body_paragraphs(body: ET.Element) -> Iterator[tuple[ET.Element, str]]:
    """按 XML 顺序遍历 body 里所有 <w:p>，包含表格 cell 内的。

    yield (p_elem, source_path)，source_path 如 "body" 或 "body > tbl[0] > tc[0,1]"。
    """
    for child in body:
        tag = child.tag.removeprefix(W)
        if tag == "p":
            yield child, "body"
        elif tag == "tbl":
            yield from _iter_tbl_paragraphs(child, "body > tbl")


def _iter_tbl_paragraphs(tbl: ET.Element, path: str) -> Iterator[tuple[ET.Element, str]]:
    rows = tbl.findall(f"{W}tr")
    for ri, tr in enumerate(rows):
        for ci, tc in enumerate(tr.findall(f"{W}tc")):
            sub_path = f"{path} > tc[{ri},{ci}]"
            for child in tc:
                ctag = child.tag.removeprefix(W)
                if ctag == "p":
                    yield child, sub_path
                elif ctag == "tbl":
                    yield from _iter_tbl_paragraphs(child, f"{sub_path} > tbl")


# -------------- 段落分类 --------------

def classify_paragraph(p: ET.Element) -> tuple[str, dict]:
    """返回 (kind, info)，kind ∈ {empty, cover, toc, heading, body}"""
    ppr = _get(p, "pPr")
    pstyle_val = _attr(_get(ppr, "pStyle") if ppr is not None else None, "val")
    jc_val = _attr(_get(ppr, "jc") if ppr is not None else None, "val")

    texts = [t.text or "" for t in p.iter(f"{W}t")]
    text = "".join(texts).strip()

    has_field = (
        p.find(f".//{W}fldChar") is not None
        or p.find(f".//{W}instrText") is not None
    )

    info = {
        "text": text,
        "pStyle": pstyle_val,
        "jc": jc_val,
        "has_field": has_field,
    }

    if not text and not has_field:
        return "empty", info

    # TOC 优先判断（字段代码 + pStyle）
    if has_field:
        for instr in p.iter(f"{W}instrText"):
            if instr.text and "TOC" in instr.text.upper():
                return "toc", info
    if pstyle_val:
        low = pstyle_val.lower()
        if low.startswith("toc") or "toc" in low or "目录" in pstyle_val or "目次" in pstyle_val:
            return "toc", info

    # 标题
    if pstyle_val:
        low = pstyle_val.lower()
        if "heading" in low or "标题" in pstyle_val or "title" in low:
            return "heading", info

    # 封面判定（综合多个弱信号）
    cover_signals = 0
    if jc_val in ("right", "center"):
        cover_signals += 1
    if has_field:
        cover_signals += 1
    if any(kw in text for kw in _COVER_FIELD_KEYWORDS):
        cover_signals += 2
    # 超短且无完整标点（不是完整句子）
    if len(text) < 12 and not any(c in text for c in "。，,、；;：:"):
        cover_signals += 1
    if cover_signals >= 2:
        return "cover", info

    return "body", info


# -------------- 从段落文字里提取老师的格式要求 --------------

# 模式：匹配老师用中文描述的格式约束。命中即提示 Agent「这里有显式要求，比 XML 推断优先级高」
_FORMAT_TEXT_PATTERNS = [
    ("font",  re.compile(r"(宋体|仿宋|黑体|楷体|微软雅黑|等线|Times\s*New\s*Roman|Arial|Calibri|Consolas)")),
    ("size",  re.compile(r"(小[一二三四五六七八]|[一二三四五六七八]号|[0-9]{1,2}\s*(?:pt|磅|号字?))")),
    ("line",  re.compile(r"(\d+(?:\.\d+)?)\s*倍(?:行距)?|行距\s*(\d+(?:\.\d+)?)")),
    ("jc",    re.compile(r"(左对齐|右对齐|居中|两端对齐|分散对齐)")),
    ("style", re.compile(r"(加粗|粗体|斜体|下划线|首行缩进|顶格)")),
]


def extract_text_format_hints(paragraph_infos: list[dict]) -> list[dict]:
    """扫段落文字，摘出含格式描述关键词的段"""
    hints = []
    for info in paragraph_infos:
        text = info.get("text") or ""
        if not text:
            continue
        matches = []
        for label, pat in _FORMAT_TEXT_PATTERNS:
            for m in pat.finditer(text):
                # 取第一个非空捕获组（或整体匹配）
                val = next((g for g in m.groups() if g), m.group(0))
                matches.append((label, val))
        if matches:
            hints.append({
                "idx": info.get("idx"),
                "path": info.get("path"),
                "text": text,
                "matches": matches,
            })
    return hints


def scan_document(doc_root: ET.Element) -> dict:
    """扫全文，分类所有段落。返回 dict 含 body/cover/toc/heading/empty 列表"""
    buckets: dict[str, list[dict]] = {k: [] for k in ("body", "cover", "toc", "heading", "empty")}
    if doc_root is None:
        return buckets
    body = doc_root.find(f"{W}body")
    if body is None:
        return buckets

    for idx, (p, path) in enumerate(iter_body_paragraphs(body)):
        kind, info = classify_paragraph(p)
        info["idx"] = idx
        info["path"] = path
        info["elem"] = p
        buckets[kind].append(info)
    return buckets


# -------------- 从正文首段抽样式 --------------

def extract_paragraph_style(p: ET.Element) -> dict:
    out = {}
    ppr = _get(p, "pPr")
    if ppr is not None:
        pstyle = _get(ppr, "pStyle")
        out["p.pStyle"] = _attr(pstyle, "val")
        spacing = _get(ppr, "spacing")
        if spacing is not None:
            out["p.line"] = spacing.get(f"{W}line")
            out["p.lineRule"] = spacing.get(f"{W}lineRule")
        ind = _get(ppr, "ind")
        if ind is not None:
            out["p.firstLineChars"] = ind.get(f"{W}firstLineChars")
        jc = _get(ppr, "jc")
        if jc is not None:
            out["p.jc"] = jc.get(f"{W}val")
        rpr_in_ppr = _get(ppr, "rPr")
        if rpr_in_ppr is not None:
            sz = _get(rpr_in_ppr, "sz")
            if sz is not None:
                out["p.size"] = sz.get(f"{W}val")
    r = _get(p, "r")
    if r is not None:
        rpr = _get(r, "rPr")
        if rpr is not None:
            sz = _get(rpr, "sz")
            if sz is not None:
                out["r.size"] = sz.get(f"{W}val")
            fonts = _get(rpr, "rFonts")
            if fonts is not None:
                out["r.ascii"] = fonts.get(f"{W}ascii")
                out["r.eastAsia"] = fonts.get(f"{W}eastAsia")
    return out


# -------------- diff 打印 --------------

def _coalesce(*vals):
    for v in vals:
        if v not in (None, ""):
            return v
    return None


def print_style_diff(styles_data: dict, body_style: dict, body_preview: str):
    print(f"  取样来源：分类为 body 的第一段（共扫出 {body_preview or '（无正文段）'} 字符）")
    if not body_style:
        print("  ⚠️  文档里没找到 body 段（可能只有封面 + 表格占位符，需用户提供正文样例）")
        return

    actual_body_size = _coalesce(
        body_style.get("r.size"), body_style.get("p.size"),
        styles_data.get("normal.size"), styles_data.get("docDefault.size"),
    )
    actual_body_line = _coalesce(
        body_style.get("p.line"),
        styles_data.get("normal.line"), styles_data.get("docDefault.line"),
    )
    actual_body_line_rule = _coalesce(
        body_style.get("p.lineRule"),
        styles_data.get("normal.lineRule"), styles_data.get("docDefault.lineRule"),
    )
    actual_body_first = _coalesce(
        body_style.get("p.firstLineChars"),
        styles_data.get("normal.firstLineChars"), styles_data.get("docDefault.firstLineChars"),
    )
    actual_body_jc = _coalesce(
        body_style.get("p.jc"), styles_data.get("normal.jc"),
    )
    actual_east_asia = _coalesce(
        body_style.get("r.eastAsia"),
        styles_data.get("normal.eastAsia"), styles_data.get("docDefault.eastAsia"),
    )
    actual_ascii = _coalesce(
        body_style.get("r.ascii"),
        styles_data.get("normal.ascii"), styles_data.get("docDefault.ascii"),
    )

    rows = [
        ("body.size",         actual_body_size,      SKILL_DEFAULTS["body.size"]),
        ("body.line",         actual_body_line,      SKILL_DEFAULTS["body.line"]),
        ("body.lineRule",     actual_body_line_rule, SKILL_DEFAULTS["body.lineRule"]),
        ("body.firstLine",    actual_body_first,     SKILL_DEFAULTS["body.firstLine"]),
        ("body.jc",           actual_body_jc,        SKILL_DEFAULTS["body.jc"]),
        ("body.fontEastAsia", actual_east_asia,      {"val": "继承模板", "note": "宋体/等线等"}),
        ("body.fontAscii",    actual_ascii,          {"val": "继承模板", "note": "Times/Calibri 等"}),
    ]

    print(f"\n  {'维度':<20} {'模板实际':<20} {'skill 默认':<20} {'是否一致':<10} 备注")
    print(f"  {'-'*20} {'-'*20} {'-'*20} {'-'*10} {'-'*30}")

    diffs = 0
    hints = []
    for name, actual, default in rows:
        default_val = default["val"]
        def_str = str(default_val)
        actual_str = str(actual) if actual is not None else "（未显式）"

        if def_str == "继承模板":
            same = True
        else:
            same = (actual_str == def_str)

        # 特殊：body.jc 中 left / both / 未显式 都算可接受
        if name == "body.jc" and actual_str in ("left", "both", "（未显式）"):
            same = True

        mark = "✓" if same else "❌ DIFF"
        if not same:
            diffs += 1
            if name == "body.size":
                hints.append(f"build_body_paragraph(text, size={actual_str}, ...)")
            elif name == "body.line":
                hints.append(f"build_body_paragraph(text, line={actual_str}, ...)")
            elif name == "body.firstLine":
                hints.append(f"build_body_paragraph(text, first_line_chars={actual_str}, ...)")
            elif name == "body.jc":
                hints.append(f"build_body_paragraph(text, jc={actual_str!r}, ...)")

        note = default["note"]
        print(f"  {name:<20} {actual_str:<20} {def_str:<20} {mark:<10} {note}")

    print()
    if diffs == 0:
        print("  ✅ 没有 diff，可直接用 builders.py 默认参数")
    else:
        print(f"  ⚠️  共 {diffs} 处 diff。写正文时给 build_body_paragraph() 传 override：")
        for h in hints:
            print(f"         {h}")


# -------------- main --------------

def main():
    if len(sys.argv) < 2:
        sys.exit("用法: inspect_template.py <docx 路径> [--verbose]")
    path = Path(sys.argv[1])
    if not path.exists():
        sys.exit(f"不存在: {path}")
    verbose = "--verbose" in sys.argv

    from docx import Document
    doc = Document(str(path))

    print(f"=== {path.name} ===\n")

    print("--- 顶层段落 ---")
    dump_paragraphs(doc.paragraphs)

    print("\n--- 表格（递归展开嵌套）---")
    if not doc.tables:
        print("  （此模板无表格——常见于论文/技术文档类模板）")
    else:
        for ti, tbl in enumerate(doc.tables):
            dump_table(tbl, str(ti))

    print("\n--- sectPr 页面 ---")
    for section in doc.sections:
        print(f"  页面大小: {section.page_width} x {section.page_height}")
        print(f"  页边距:   上 {section.top_margin}  下 {section.bottom_margin}  "
              f"左 {section.left_margin}  右 {section.right_margin}")

    # 段落分类
    with ZipFile(str(path)) as zf:
        styles_root = _read_xml(zf, "word/styles.xml")
        doc_root = _read_xml(zf, "word/document.xml")

    buckets = scan_document(doc_root)

    print("\n--- 段落分类（按文档顺序扫描） ---")
    print(f"  封面 / 元信息：{len(buckets['cover'])} 段")
    print(f"  目录（TOC）    ：{len(buckets['toc'])} 段")
    print(f"  标题          ：{len(buckets['heading'])} 段")
    print(f"  正文          ：{len(buckets['body'])} 段")
    print(f"  空段          ：{len(buckets['empty'])} 段")

    # ★ 封面保留区警告
    if buckets["cover"]:
        print("\n--- 🛡️  封面 / 元信息（保留区，拼接阶段不要动）---")
        print("  这些段常用表单字段、文本框、绝对定位等「魔法」排版，覆盖/删除/重写会破坏布局。")
        print("  拼接 docx 时应原样保留这些段的 XML，只在其后追加正文。")
        for item in buckets["cover"][:15]:
            text = item["text"][:60] + ("…" if len(item["text"]) > 60 else "")
            reasons = []
            if item["jc"] in ("right", "center"):
                reasons.append(f"jc={item['jc']}")
            if item["has_field"]:
                reasons.append("fldChar")
            if any(kw in item["text"] for kw in _COVER_FIELD_KEYWORDS):
                reasons.append("关键词")
            reason_str = "/".join(reasons) if reasons else "启发式"
            print(f"  [#{item['idx']:3d} {item['path']}] ({reason_str}) {text!r}")
        if len(buckets["cover"]) > 15:
            print(f"  … 还有 {len(buckets['cover']) - 15} 段未列出")

    if buckets["toc"]:
        print(f"\n--- 📑 目录（TOC）---")
        print("  若报告重写目录，Word 打开后 F9 刷新即可；否则保留原段。")
        for item in buckets["toc"][:10]:
            text = item["text"][:60]
            print(f"  [#{item['idx']:3d} {item['path']}] {text!r}")

    # ★ 文字格式要求（老师显式说明的优先级高于 XML 推断）
    all_text_paras = (
        buckets["body"] + buckets["heading"] + buckets["cover"] + buckets["toc"]
    )
    text_hints = extract_text_format_hints(all_text_paras)
    if text_hints:
        print("\n--- 📜 模板内的文字格式要求（老师显式说明，优先级高于 XML 推断）---")
        print("  发现以下段落含格式关键字。若与 XML 样式不符，以文字要求为准。")
        for h in text_hints[:8]:
            kw_str = " / ".join(f"{l}={v}" for l, v in h["matches"])
            text_short = h["text"][:100] + ("…" if len(h["text"]) > 100 else "")
            print(f"  [#{h['idx']:3d} {h['path']}]")
            print(f"    命中: {kw_str}")
            print(f"    原文: {text_short!r}")
        if len(text_hints) > 8:
            print(f"  … 还有 {len(text_hints) - 8} 段命中未列出")

    # ★ 样式 diff —— 取首个 body 段
    print("\n--- 样式 diff（模板正文首段 vs skill 硬编码默认）---")
    styles_data = extract_default_styles(styles_root)

    body_items = buckets["body"]
    if body_items:
        first_body = body_items[0]
        body_preview = first_body["text"][:60]
        print(f"  正文首段位置：#{first_body['idx']} @ {first_body['path']}")
        print(f"  首段预览：{body_preview!r}")
        body_style = extract_paragraph_style(first_body["elem"])
        print_style_diff(styles_data, body_style, body_preview)
    else:
        print("  ⚠️  没找到正文段：模板可能只有封面 + TOC + 空占位 cell")
        print("     建议让用户提供一份填了正文的样例，或直接进入阶段 2（md 成稿）再决定样式。")

    if verbose:
        print("\n--- raw styles.xml 提取值 ---")
        for k, v in styles_data.items():
            print(f"  {k}: {v}")
        if body_items:
            print("\n--- raw first body 段 style ---")
            for k, v in extract_paragraph_style(body_items[0]["elem"]).items():
                print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
